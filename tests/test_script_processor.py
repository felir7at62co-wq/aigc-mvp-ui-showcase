import json
import importlib.util
import os
import tempfile
import unittest

from docx import Document

from agents.script_processor import ScriptProcessor, normalize_episode_number
from core.project import Project


class ScriptProcessorTests(unittest.TestCase):
    def _write_text(self, directory, content, encoding="utf-8"):
        path = os.path.join(directory, "script.txt")
        with open(path, "w", encoding=encoding) as handle:
            handle.write(content)
        return path

    def test_normalizes_arabic_and_chinese_numbers(self):
        self.assertEqual(normalize_episode_number("第3集"), "03")
        self.assertEqual(normalize_episode_number("第一百二十三集"), "123")

    def test_splits_only_line_start_markers_and_filters_high_confidence_lines(self):
        with tempfile.TemporaryDirectory() as directory:
            source = self._write_text(
                directory,
                "\n".join([
                    "第1集：开始",
                    "版本二：旁白：故事开始。",
                    "角色：第2集的时候我再告诉你。",
                    "【下集钩子】保留这句正文",
                    "第2集 继续",
                    "旁白：第二集正文。",
                    "下集预告：",
                    "这段预告应删除。",
                ]),
            )
            analysis = ScriptProcessor().analyze(source)

        self.assertEqual([episode.number for episode in analysis.episodes], ["01", "02"])
        self.assertIn("旁白：故事开始。", analysis.episodes[0].cleaned_content)
        self.assertIn("角色：第2集的时候我再告诉你。", analysis.episodes[0].cleaned_content)
        self.assertIn("保留这句正文", analysis.episodes[0].cleaned_content)
        self.assertNotIn("预告应删除", analysis.episodes[1].cleaned_content)

    def test_accepts_bracketed_episode_titles(self):
        with tempfile.TemporaryDirectory() as directory:
            source = self._write_text(
                directory,
                "\n".join([
                    "第1集",
                    "第一集正文",
                    "【第2集：新的标题】",
                    "第二集正文",
                    "## 第3集 - 结尾",
                    "第三集正文",
                ]),
            )
            analysis = ScriptProcessor().analyze(source)

        self.assertEqual(
            [episode.number for episode in analysis.episodes],
            ["01", "02", "03"],
        )
        self.assertFalse(analysis.warnings)

    def test_reports_gaps_duplicates_and_single_episode_fallback(self):
        with tempfile.TemporaryDirectory() as directory:
            gap_source = self._write_text(directory, "第1集\n正文\n第3集\n正文")
            gap = ScriptProcessor().analyze(gap_source)
            self.assertTrue(any("跳到" in warning for warning in gap.warnings))

            duplicate_source = self._write_text(directory, "第1集\n甲\n第1集\n乙")
            duplicate = ScriptProcessor().analyze(duplicate_source)
            self.assertTrue(duplicate.errors)

            single_source = self._write_text(directory, "没有集数标题的完整正文")
            single = ScriptProcessor().analyze(single_source)
            self.assertEqual(single.episodes[0].number, "01")
            self.assertTrue(single.warnings)

    def test_word_weak_signals_are_suspect_by_default(self):
        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "script.docx")
            document = Document()
            document.add_heading("第1集：标题", level=1)
            document.add_heading("短标题", level=2)
            document.add_paragraph("角色：对白正文")
            document.save(path)

            analysis = ScriptProcessor().analyze(path)

        suspects = [
            decision for decision in analysis.episodes[0].decisions
            if decision.action == "suspect"
        ]
        self.assertTrue(any("Word 标题样式" in decision.reason for decision in suspects))
        self.assertIn("短标题", analysis.episodes[0].cleaned_content)

    def test_txt_gbk_and_unsupported_doc(self):
        with tempfile.TemporaryDirectory() as directory:
            source = self._write_text(directory, "第1集\n中文正文", encoding="gbk")
            analysis = ScriptProcessor().analyze(source)
            self.assertIn("中文正文", analysis.episodes[0].cleaned_content)

            doc_path = os.path.join(directory, "legacy.doc")
            with open(doc_path, "wb") as handle:
                handle.write(b"legacy")
            with self.assertRaisesRegex(ValueError, "仅支持"):
                ScriptProcessor().analyze(doc_path)

    def test_text_quality_reports_duplicate_and_long_lines(self):
        with tempfile.TemporaryDirectory() as directory:
            repeated = "这是一段需要检查的重复正文内容"
            source = self._write_text(
                directory,
                f"第1集\n{repeated}\n{repeated}\n{'长' * 301}",
            )
            analysis = ScriptProcessor().analyze(source)
            warnings = analysis.episodes[0].warnings
            self.assertTrue(any("重复" in warning for warning in warnings))
            self.assertTrue(any("300" in warning for warning in warnings))

    def test_save_writes_manifest_and_rejects_duplicates(self):
        with tempfile.TemporaryDirectory() as directory:
            source = self._write_text(directory, "第1集\n正文一\n第2集\n正文二")
            project = Project.create(directory, "demo", source)
            processor = ScriptProcessor()
            analysis = processor.analyze(source)
            result = processor.save(project.project_dir, analysis)

            self.assertTrue(result.success, result.error)
            with open(result.manifest_path, "r", encoding="utf-8") as handle:
                manifest = json.load(handle)
            self.assertEqual(set(manifest["episodes"]), {"01", "02"})
            self.assertEqual(len(manifest["episodes"]["01"]["sha256"]), 64)

            stale_path = os.path.join(project.project_dir, "episodes", "03.txt")
            with open(stale_path, "w", encoding="utf-8") as handle:
                handle.write("旧分集")
            second_save = processor.save(project.project_dir, analysis)
            self.assertTrue(second_save.success)
            self.assertFalse(os.path.exists(stale_path))

            duplicate_source = self._write_text(directory, "第1集\n甲\n第1集\n乙")
            duplicate = processor.analyze(duplicate_source)
            failed = processor.save(project.project_dir, duplicate)
            self.assertFalse(failed.success)

    def test_config_save_preserves_unknown_fields_and_drops_obsolete_generation_sections(self):
        if importlib.util.find_spec("yaml") is None:
            self.skipTest("PyYAML is not installed")
        from core.config import load_config, save_config

        with tempfile.TemporaryDirectory() as directory:
            config_path = os.path.join(directory, "config.yaml")
            with open(config_path, "w", encoding="utf-8") as handle:
                handle.write(
                    "custom_section:\n  retained: true\n"
                    "storyboard:\n  output_format: jpg\n"
                )
            config = load_config(config_path, base_dir=directory)
            save_config(config, config_path)
            with open(config_path, "r", encoding="utf-8") as handle:
                saved = handle.read()

            self.assertIn("custom_section:", saved)
            self.assertNotIn("storyboard:", saved)


if __name__ == "__main__":
    unittest.main()
