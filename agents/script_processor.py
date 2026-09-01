"""Deterministic screenplay splitting, filtering, preview, and persistence."""

from __future__ import annotations

import hashlib
import json
import os
import re
import tempfile
from dataclasses import asdict, dataclass, field
from typing import Any, Dict, List, Mapping, Optional, Sequence, Tuple

from docx import Document


EPISODE_MANIFEST = "manifest.json"


@dataclass
class FilterDecision:
    line_number: int
    original: str
    processed: str
    action: str  # removed / modified / suspect
    reason: str
    restored: bool = False


@dataclass
class EpisodeAnalysis:
    number: str
    marker: str
    original_content: str
    cleaned_content: str
    source_start_line: int = 1
    decisions: List[FilterDecision] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    content_hash: str = ""

    def refresh_hash(self) -> str:
        self.content_hash = content_sha256(self.cleaned_content)
        return self.content_hash


@dataclass
class ScriptAnalysis:
    source_path: str
    original_text: str
    episodes: List[EpisodeAnalysis]
    filter_options: Dict[str, bool]
    warnings: List[str] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)


@dataclass
class SaveResult:
    success: bool
    episode_count: int = 0
    manifest_path: str = ""
    error: str = ""


DEFAULT_FILTER_OPTIONS: Dict[str, bool] = {
    "filter_episode_title": True,
    "filter_version": True,
    "filter_next_episode": True,
    # Weak signals are review-only unless explicitly enabled.
    "filter_colon_title": False,
    "filter_word_heading": False,
    "filter_large_font": False,
    "filter_short_bold": False,
    "filter_bracket_title": False,
    "filter_table_content": False,
}


_CHINESE_DIGITS = {
    "零": 0, "〇": 0, "一": 1, "二": 2, "两": 2, "三": 3,
    "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9,
}
_CHINESE_UNITS = {"十": 10, "百": 100, "千": 1000}

_CN_MARKER_RE = re.compile(
    r"^\s*第\s*(?P<number>[零〇一二两三四五六七八九十百千\d]+)\s*"
    r"[集话回章](?P<title>\s*$|\s+[^\r\n]+$|\s*[：:\-—]\s*.*$)",
    re.IGNORECASE,
)
_EN_MARKER_RE = re.compile(
    r"^\s*(?:EP|EPISODE)\s*[._\-]?\s*(?P<number>\d+)"
    r"(?P<title>\s*$|\s+[^\r\n]+$|\s*[：:\-—]\s*.*$)",
    re.IGNORECASE,
)


def content_sha256(content: str) -> str:
    return hashlib.sha256(content.encode("utf-8")).hexdigest()


def chinese_number_to_int(value: str) -> int:
    value = value.strip()
    if value.isdigit():
        return int(value)
    if not value:
        raise ValueError("集数为空")

    total = 0
    number = 0
    for char in value:
        if char in _CHINESE_DIGITS:
            number = _CHINESE_DIGITS[char]
        elif char in _CHINESE_UNITS:
            unit = _CHINESE_UNITS[char]
            total += (number or 1) * unit
            number = 0
        else:
            raise ValueError(f"无法识别中文集数: {value}")
    total += number
    if total <= 0:
        raise ValueError(f"无效集数: {value}")
    return total


def normalize_episode_number(episode_str: str) -> str:
    match = re.search(r"\d+", episode_str)
    if match:
        number = int(match.group())
    else:
        match = re.search(r"[零〇一二两三四五六七八九十百千]+", episode_str)
        if not match:
            raise ValueError(f"无法识别集数: {episode_str}")
        number = chinese_number_to_int(match.group())
    if number < 1 or number > 9999:
        raise ValueError(f"集数超出范围: {number}")
    return f"{number:02d}"


def _read_text_file(file_path: str) -> str:
    last_error: Optional[Exception] = None
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            with open(file_path, "r", encoding=encoding) as handle:
                return handle.read()
        except UnicodeDecodeError as exc:
            last_error = exc
    raise ValueError(f"无法识别文本编码: {last_error}")


def read_script(file_path: str) -> Dict[str, Any]:
    """Read a DOCX or TXT file and retain line-level Word formatting metadata."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt":
        return {"text": _read_text_file(file_path), "format_info": []}
    if ext != ".docx":
        raise ValueError(f"不支持的文件格式: {ext or '无扩展名'}（仅支持 .docx / .txt）")

    document = Document(file_path)
    lines: List[str] = []
    format_info: List[Dict[str, Any]] = []

    def append_line(text: str, **metadata: Any) -> None:
        cleaned = text.strip()
        if not cleaned:
            return
        if lines:
            lines.append("")
        lines.append(cleaned)
        format_info.append({"line_number": len(lines), "text": cleaned, **metadata})

    for paragraph in document.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        sizes = [run.font.size.pt for run in paragraph.runs if run.font.size]
        append_line(
            paragraph.text,
            is_title=style_name.startswith("Heading") or "标题" in style_name,
            style=style_name,
            font_size=round(max(sizes)) if sizes else None,
            bold=any(run.bold is True for run in paragraph.runs),
            italic=any(run.italic is True for run in paragraph.runs),
            is_table=False,
        )

    for table in document.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            append_line(
                row_text,
                is_title=False,
                style="Table Content",
                font_size=None,
                bold=False,
                italic=False,
                is_table=True,
            )

    return {"text": "\n".join(lines), "format_info": format_info}


def _episode_marker(line: str) -> Optional[Tuple[str, str]]:
    candidate = line.strip()
    candidate = re.sub(r"^[#>*-]+\s*", "", candidate)
    candidate = re.sub(r"^[\u3010\[\(\uff08]\s*", "", candidate)
    candidate = re.sub(r"[\u3011\]\)\uff09]\s*$", "", candidate)
    for pattern in (_CN_MARKER_RE, _EN_MARKER_RE):
        match = pattern.match(candidate)
        if match:
            return normalize_episode_number(match.group("number")), line.strip()
    return None


def _format_lookup(format_info: Sequence[Mapping[str, Any]]) -> Dict[int, Mapping[str, Any]]:
    return {
        int(item["line_number"]): item
        for item in format_info
        if item.get("line_number") is not None
    }


def _weak_reason(line: str, metadata: Mapping[str, Any]) -> Optional[Tuple[str, str]]:
    if metadata.get("is_table"):
        return "filter_table_content", "Word 表格内容"
    if metadata.get("is_title"):
        return "filter_word_heading", "Word 标题样式"
    font_size = metadata.get("font_size")
    if isinstance(font_size, (int, float)) and font_size >= 14:
        return "filter_large_font", f"大字号（{font_size:g} 磅）"
    if metadata.get("bold") and len(line) <= 20:
        return "filter_short_bold", "短粗体行"
    if re.match(r"^[：:].{1,10}$", line):
        return "filter_colon_title", "冒号开头短标题"
    if re.match(r"^(?:【.{1,15}】|\[.{1,15}\]|\(.{1,15}\)|\{.{1,15}\})$", line):
        return "filter_bracket_title", "短括号标题"
    return None


def _clean_episode_lines(
    lines: Sequence[Tuple[int, str]],
    marker_line: Optional[int],
    metadata_by_line: Mapping[int, Mapping[str, Any]],
    options: Mapping[str, bool],
) -> Tuple[str, List[FilterDecision]]:
    output: List[str] = []
    decisions: List[FilterDecision] = []
    truncate_preview = False

    for line_number, raw_line in lines:
        line = raw_line.strip()
        if marker_line == line_number and options.get("filter_episode_title", True):
            decisions.append(FilterDecision(line_number, raw_line, "", "removed", "集数标题"))
            continue
        if truncate_preview:
            if line:
                decisions.append(FilterDecision(line_number, raw_line, "", "removed", "下集预告正文"))
            continue
        if not line:
            if output and output[-1] != "":
                output.append("")
            continue

        if options.get("filter_next_episode", True) and re.match(
            r"^(?:【\s*)?(?:下集预告|下一集预告|预告片)(?:\s*】)?(?:\s*[：:].*)?$",
            line,
        ):
            decisions.append(FilterDecision(line_number, raw_line, "", "removed", "下集预告标记"))
            truncate_preview = True
            continue

        processed = line
        if options.get("filter_version", True):
            version_match = re.match(
                r"^版本\s*[零〇一二两三四五六七八九十百千\d]+\s*[：:]\s*(.*)$",
                processed,
            )
            if version_match:
                remainder = version_match.group(1).strip()
                action = "modified" if remainder else "removed"
                decisions.append(FilterDecision(line_number, raw_line, remainder, action, "版本前缀"))
                if not remainder:
                    continue
                processed = remainder

        if options.get("filter_next_episode", True) and "【下集钩子】" in processed:
            remainder = processed.replace("【下集钩子】", "").strip()
            action = "modified" if remainder else "removed"
            decisions.append(FilterDecision(line_number, raw_line, remainder, action, "下集钩子标签"))
            if not remainder:
                continue
            processed = remainder

        weak = _weak_reason(processed, metadata_by_line.get(line_number, {}))
        if weak:
            option_name, reason = weak
            if options.get(option_name, False):
                decisions.append(FilterDecision(line_number, raw_line, "", "removed", reason))
                continue
            decisions.append(FilterDecision(line_number, raw_line, processed, "suspect", reason))

        output.append(processed)

    while output and output[-1] == "":
        output.pop()
    return "\n".join(output).strip(), decisions


class ScriptProcessor:
    def __init__(self, filter_options: Optional[Mapping[str, bool]] = None):
        self.filter_options = dict(DEFAULT_FILTER_OPTIONS)
        if filter_options:
            self.filter_options.update({key: bool(value) for key, value in filter_options.items()})

    def analyze(self, script_path: str) -> ScriptAnalysis:
        script_data = read_script(script_path)
        text = script_data["text"]
        lines = text.splitlines()
        format_by_line = _format_lookup(script_data.get("format_info", []))

        markers: List[Tuple[int, str, str]] = []
        for index, line in enumerate(lines, start=1):
            marker = _episode_marker(line)
            if marker:
                markers.append((index, marker[0], marker[1]))

        warnings: List[str] = []
        errors: List[str] = []
        if not markers:
            warnings.append("未识别到集数标记，已按单集 01 处理")
            markers = [(1, "01", "01")]
            synthetic_marker = True
        else:
            synthetic_marker = False

        numbers = [int(number) for _, number, _ in markers]
        duplicate_numbers = sorted({number for number in numbers if numbers.count(number) > 1})
        if duplicate_numbers:
            errors.append("检测到重复集号: " + ", ".join(f"{number:02d}" for number in duplicate_numbers))
        for previous, current in zip(numbers, numbers[1:]):
            if current > previous + 1:
                warnings.append(f"集号从 {previous:02d} 跳到 {current:02d}")
            elif current < previous:
                warnings.append(f"集号顺序异常: {previous:02d} 后出现 {current:02d}")

        episodes: List[EpisodeAnalysis] = []
        for marker_index, (line_number, number, marker_text) in enumerate(markers):
            start_index = 0 if synthetic_marker else line_number - 1
            end_index = markers[marker_index + 1][0] - 1 if marker_index + 1 < len(markers) else len(lines)
            episode_lines = [(index + 1, lines[index]) for index in range(start_index, end_index)]
            marker_line = None if synthetic_marker else line_number
            cleaned, decisions = _clean_episode_lines(
                episode_lines,
                marker_line,
                format_by_line,
                self.filter_options,
            )
            episode_warnings: List[str] = []
            if not cleaned:
                episode_warnings.append("过滤后正文为空")
            episode_warnings.extend(_text_quality_warnings(cleaned))
            original = "\n".join(line for _, line in episode_lines).strip()
            episode = EpisodeAnalysis(
                number=number,
                marker=marker_text,
                original_content=original,
                cleaned_content=cleaned,
                source_start_line=start_index + 1,
                decisions=decisions,
                warnings=episode_warnings,
            )
            episode.refresh_hash()
            episodes.append(episode)

        return ScriptAnalysis(
            source_path=os.path.abspath(script_path),
            original_text=text,
            episodes=episodes,
            filter_options=dict(self.filter_options),
            warnings=warnings,
            errors=errors,
        )

    def save(self, project_dir: str, analysis: ScriptAnalysis) -> SaveResult:
        try:
            if analysis.errors:
                raise ValueError("；".join(analysis.errors))
            numbers = [episode.number for episode in analysis.episodes]
            if len(numbers) != len(set(numbers)):
                raise ValueError("存在重复集号，不能保存")
            if not analysis.episodes:
                raise ValueError("没有可保存的分集")
            for episode in analysis.episodes:
                episode.cleaned_content = episode.cleaned_content.strip()
                if not episode.cleaned_content:
                    raise ValueError(f"第{episode.number}集正文为空")
                episode.refresh_hash()

            episodes_dir = os.path.join(project_dir, "episodes")
            os.makedirs(episodes_dir, exist_ok=True)
            staged: Dict[str, str] = {}
            try:
                for episode in analysis.episodes:
                    with tempfile.NamedTemporaryFile(
                        "w", encoding="utf-8", delete=False, dir=episodes_dir, suffix=".tmp"
                    ) as handle:
                        handle.write(episode.cleaned_content)
                        staged[episode.number] = handle.name

                desired = {f"{number}.txt" for number in numbers}
                for filename in os.listdir(episodes_dir):
                    if filename.endswith(".txt") and filename not in desired:
                        os.remove(os.path.join(episodes_dir, filename))
                for number, temp_path in staged.items():
                    os.replace(temp_path, os.path.join(episodes_dir, f"{number}.txt"))
                staged.clear()
            finally:
                for temp_path in staged.values():
                    if os.path.exists(temp_path):
                        os.remove(temp_path)

            source_mtime = os.path.getmtime(analysis.source_path) if os.path.exists(analysis.source_path) else 0
            manifest = {
                "version": 1,
                "source_path": analysis.source_path,
                "source_mtime": source_mtime,
                "filter_options": analysis.filter_options,
                "episodes": {
                    episode.number: {
                        "sha256": episode.content_hash,
                        "char_count": len(episode.cleaned_content),
                    }
                    for episode in analysis.episodes
                },
            }
            manifest_path = os.path.join(episodes_dir, EPISODE_MANIFEST)
            _write_json_atomic(manifest_path, manifest)

            from core.project import Project

            project = Project.load(project_dir)
            project.state.episode_count = len(analysis.episodes)
            project.sync_episode_hashes({
                episode.number: episode.content_hash for episode in analysis.episodes
            })
            project.mark_step_completed(
                "script",
                len(analysis.episodes),
                {"manifest": manifest_path, "source_path": analysis.source_path},
            )
            return SaveResult(True, len(analysis.episodes), manifest_path)
        except Exception as exc:
            return SaveResult(False, error=str(exc))

    def process(
        self,
        script_path: str,
        output_dir: str,
        start_episode: int = 1,
        end_episode: int = 999,
    ) -> Dict[str, Any]:
        """Compatibility wrapper for non-UI callers."""
        analysis = self.analyze(script_path)
        selected = [
            episode for episode in analysis.episodes
            if start_episode <= int(episode.number) <= end_episode
        ]
        if analysis.errors:
            return {"success": False, "error": "；".join(analysis.errors), "episodes": []}
        os.makedirs(output_dir, exist_ok=True)
        results = []
        for episode in selected:
            path = os.path.join(output_dir, f"{episode.number}.txt")
            with open(path, "w", encoding="utf-8") as handle:
                handle.write(episode.cleaned_content)
            results.append({
                "number": episode.number,
                "marker": episode.marker,
                "path": path,
                "char_count": len(episode.cleaned_content),
            })
        return {"success": True, "episode_count": len(results), "episodes": results}


def _text_quality_warnings(content: str) -> List[str]:
    """Fast deterministic checks that help creators focus manual review."""
    if not content:
        return []
    warnings: List[str] = []
    compact_length = len(re.sub(r"\s+", "", content))
    if compact_length < 20:
        warnings.append("正文过短，请确认是否拆分错误")
    elif compact_length > 12000:
        warnings.append("正文过长，可能影响镜头脚本和视频生成质量")
    lines = [line.strip() for line in content.splitlines() if line.strip()]
    if any(len(line) > 300 for line in lines):
        warnings.append("存在超过 300 字的长段落")
    seen = set()
    if any(line in seen or seen.add(line) for line in lines if len(line) >= 12):
        warnings.append("存在明显重复段落")
    if any(re.search(r"下集预告|下一集预告|预告片", line) for line in lines):
        warnings.append("正文可能仍包含下集预告")
    return warnings


def _write_json_atomic(path: str, payload: Mapping[str, Any]) -> None:
    directory = os.path.dirname(path)
    os.makedirs(directory, exist_ok=True)
    with tempfile.NamedTemporaryFile(
        "w", encoding="utf-8", delete=False, dir=directory, suffix=".tmp"
    ) as handle:
        json.dump(payload, handle, ensure_ascii=False, indent=2)
        temp_path = handle.name
    os.replace(temp_path, path)


def filter_content(
    text: str,
    format_info: Optional[List[Dict[str, Any]]] = None,
    filter_options: Optional[Mapping[str, bool]] = None,
) -> str:
    options = dict(DEFAULT_FILTER_OPTIONS)
    if filter_options:
        options.update(filter_options)
    cleaned, _ = _clean_episode_lines(
        list(enumerate(text.splitlines(), start=1)),
        None,
        _format_lookup(format_info or []),
        options,
    )
    return cleaned


def split_episodes(
    content: str,
    format_info: Optional[List[Dict[str, Any]]] = None,
    apply_filter: bool = True,
) -> List[Tuple[str, str]]:
    """Compatibility helper returning (marker, content) tuples."""
    lines = content.splitlines()
    markers = [
        (index, marker)
        for index, line in enumerate(lines)
        if (marker := _episode_marker(line)) is not None
    ]
    if not markers:
        body = filter_content(content, format_info) if apply_filter else content.strip()
        return [("01", body)]
    result: List[Tuple[str, str]] = []
    for position, (start, marker) in enumerate(markers):
        end = markers[position + 1][0] if position + 1 < len(markers) else len(lines)
        body = "\n".join(lines[start + 1:end]).strip()
        if apply_filter:
            body = filter_content(body, format_info)
        result.append((marker[1], body))
    return result


def analysis_to_dict(analysis: ScriptAnalysis) -> Dict[str, Any]:
    return asdict(analysis)
